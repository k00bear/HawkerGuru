import os
from pathlib import Path
from typing import List, Dict, Tuple, Optional, Union
from dataclasses import dataclass
from datetime import datetime
import re
import json
import docx
import streamlit as st
from dotenv import load_dotenv
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_openai import OpenAIEmbeddings, ChatOpenAI
from langchain_community.vectorstores import FAISS
from langchain.chains import ConversationalRetrievalChain
from langchain.docstore.document import Document
from langchain_community.document_loaders import TextLoader, UnstructuredWordDocumentLoader
from langchain.prompts import PromptTemplate

# Environment setup
if load_dotenv('.env'):
    OPENAI_KEY = os.getenv('OPENAI_API_KEY')
else:
    OPENAI_KEY = st.secrets['OPENAI_API_KEY']

PROJECT_ROOT = Path(__file__).parent.parent.parent
DATA_DIR = os.path.join(PROJECT_ROOT, 'data')

@dataclass
class ProcessedDocument:
    """Container for processed document content and metadata."""
    content: str
    metadata: Dict
    special_notes: Optional[Dict] = None

class BasePreprocessor:
    """Base class for document preprocessing."""
    
    def __init__(self, input_text: str):
        self.raw_text = input_text
        self.processed_text = ""
        
    def process(self) -> str:
        """Main processing pipeline to be implemented by subclasses."""
        raise NotImplementedError
    
    def get_section_content(self, section_name: str) -> str:
        """Extract content of a specific section."""
        if not self.processed_text:
            self.process()
            
        pattern = fr"## {section_name}.*?(?=## |\Z)"  # Note the 'r' prefix
        match = re.search(pattern, self.processed_text, re.DOTALL)
        return match.group(0) if match else ""

class FAQPreprocessor(BasePreprocessor):
    """Preprocesses FAQ document."""
    
    def process(self) -> str:
        """Process FAQ document."""
        text = self.raw_text
        
        # Basic cleanup first
        text = self._basic_cleanup(text)
        
        # Remove internal notes
        text = self._remove_internal_notes(text)
        
        # Remove revision history
        text = self._remove_revision_history(text)
        
        # Format main sections
        text = self._format_sections(text)
        
        # Format individual FAQs
        text = self._format_questions(text)
        
        # Add metadata
        text = self._add_metadata(text)
        
        self.processed_text = text
        return text
    
    def _basic_cleanup(self, text: str) -> str:
        """Initial cleanup of text."""
        # Remove any null bytes and normalize whitespace
        text = text.replace('\x00', '')
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        
        # Remove image references and other artifacts
        text = re.sub(r'!\[.*?\]\(.*?\)', '', text)
        text = re.sub(r'\{.*?\}', '', text)
        text = re.sub(r'\[.*?\]', '', text)
        
        # Remove empty lines and normalize spacing
        lines = [line.strip() for line in text.split('\n')]
        text = '\n'.join(line for line in lines if line)
        
        return text
    
    def _remove_internal_notes(self, text: str) -> str:
        """Remove internal CC notes."""
        # Remove anything between [[ and ]]
        text = re.sub(r'\[\[.*?\]\].*?\+=+\+', '', text, flags=re.DOTALL)
        return text
    
    def _remove_revision_history(self, text: str) -> str:
        """Remove revision history section."""
        # Remove the revision history section and everything after it
        text = re.sub(r'REVISION HISTORY.*$', '', text, flags=re.DOTALL | re.IGNORECASE)
        return text.strip()
    
    def _format_sections(self, text: str) -> str:
        """Format main FAQ sections."""
        # Start with the title
        result = "# Frequently Asked Questions (FAQs) on Electronic Tender (e-Tender)\n\n"
        
        # Define section patterns with more flexible matching
        sections = {
            r'Eligibility\s*(?:\n|$)': "Eligibility",
            r'About\s+e-?Tender\s*(?:\n|$)': "About e-Tender",
            r'Preparations?\s+Before\s+Tendering': "Preparations Before Tendering",
            r'Considerations?\s+to\s+Tender': "Considerations for Tendering",
            r'Submitting\s+a\s+Tender\s+Bid': "Submitting a Tender Bid",
            r'Amend\s+Particulars': "Amending Tender Details",
            r'Withdrawal\s+of\s+Tender': "Withdrawal of Tender",
            r'Tender\s+Results': "Tender Results",
            r'Successful\s+Tender': "Successful Tender",
            r'Protection\s+Against\s+Scams': "Protection Against Scams",
            r'Security\s+and\s+Data\s+Protection': "Security and Data Protection",
            r'Contact\s+Information': "Contact Information"
        }
        
        current_text = text
        for pattern, section_name in sections.items():
            # Look for the section pattern
            matches = re.split(f'({pattern})', current_text, flags=re.IGNORECASE)
            if len(matches) > 1:
                # Add any text before the section
                if matches[0].strip():
                    result += matches[0].strip() + "\n\n"
                # Add the section header
                result += f"## {section_name}\n\n"
                # Continue with remaining text
                current_text = ''.join(matches[2:])
        
        # Add any remaining text
        if current_text.strip():
            result += current_text.strip() + "\n"
        
        return result
    
    def _format_questions(self, text: str) -> str:
        """Format individual FAQ entries."""
        # Pattern to match numbered questions, handling various formats
        patterns = [
            # Standard numbered questions
            (r'(\d+)[\.|\)]\s*(?:\*\*)?(.*?)(?:\*\*)?\s*(?=\n|$)', r'\n### Q\1: \2\n'),
            # Questions ending with question mark
            (r'^([^#\n].*?\?)\s*$', r'\n### \1\n'),
            # Cleanup excessive newlines
            (r'\n{3,}', '\n\n'),
            # Format bullet points
            (r'(?m)^[•●]\s+', '- '),
            # Remove any remaining asterisks
            (r'\*\*', '')
        ]
        
        # Apply all patterns
        for pattern, replacement in patterns:
            text = re.sub(pattern, replacement, text, flags=re.MULTILINE)
        
        return text
    
    def _add_metadata(self, text: str) -> str:
        """Add metadata for better retrieval."""
        metadata = """<!-- Document Metadata
Type: FAQ
Purpose: Quick reference guide for e-tender process
Target Users: New hawker stall tenderers
Last Updated: March 2024
Topics: eligibility, e-tender process, payments, security, data protection
-->\n\n"""
        return metadata + text

class TenderTermsPreprocessor(BasePreprocessor):
    """Preprocesses Terms and Conditions document."""
    
    def process(self) -> str:
        """Process Terms and Conditions document."""
        text = self.raw_text
        
        # Add consistent title formatting
        text = self._format_main_title(text)
        
        # Standardize section headers
        text = self._format_section_headers(text)
        
        # Format clause numbers and subsections
        text = self._format_clauses(text)
        
        # Add section metadata
        text = self._add_section_metadata(text)
        
        self.processed_text = text
        return text
    
    def _format_main_title(self, text: str) -> str:
        """Format the main title section."""
        title_pattern = r"TERMS AND CONDITIONS OF TENDER\s*\nVer \d+: [A-Za-z]+ \d{4}"
        replacement = """
# TERMS AND CONDITIONS OF TENDER
## Version Information
{submatch}

## Document Overview
This document outlines the terms and conditions for hawker stall tenders.
"""
        return re.sub(title_pattern, lambda m: replacement.format(submatch=m.group(0)), text)
    
    def _format_section_headers(self, text: str) -> str:
        """Format all major section headers."""
        section_headers = {
            "Eligibility": 2,
            "Tendering": 2,
            "Market Stall": 2,
            "Cooked Food Stall": 2,
            "Successful Tenderer": 2,
            "Anti-Collusion": 2,
            "Reporting of Anti-competitive Conduct": 3,
            "Warranty": 3,
            "Disclosure of Prior Anti-competitive Conduct": 3
        }
        
        for header, level in section_headers.items():
            pattern = f"({header}\\s*\n)"
            hashes = "#" * level
            text = re.sub(pattern, f"\n{hashes} {header}\n\n", text)
        
        return text
    
    def _format_clauses(self, text: str) -> str:
        """Format clause numbers and subsections."""
        text = re.sub(r"^(\d+)\.\s+", r"\n### Clause \1\n", text, flags=re.MULTILINE)
        text = re.sub(r"^(\d+\.\d+)\s+", r"#### \1\n", text, flags=re.MULTILINE)
        text = re.sub(r"^(\d+\.\d+\.\d+)\s+", r"##### \1\n", text, flags=re.MULTILINE)
        return text
    
    def _add_section_metadata(self, text: str) -> str:
        """Add metadata tags for better context retrieval."""
        sections = {
            "Eligibility": "Clauses 2-6: Requirements for tender submission",
            "Tendering": "Clauses 7-17: Process and rules for submitting tenders",
            "Market Stall": "Clause 18: Specific rules for market stalls",
            "Cooked Food Stall": "Clauses 19-23: Specific rules for cooked food stalls",
            "Successful Tenderer": "Clauses 24-40: Obligations and requirements for successful tenderers",
            "Anti-Collusion": "Clause 41: Rules preventing anti-competitive behavior"
        }
        
        for section, description in sections.items():
            pattern = f"(## {section})"
            metadata = f"""
<!-- Section Metadata
Type: Terms and Conditions
Section: {section}
Description: {description}
-->

"""
            text = re.sub(pattern, f"{metadata}\\1", text)
        
        return text

class TenderNoticePreprocessor(BasePreprocessor):
    """Preprocesses Tender Notice document."""
    
    def __init__(self, input_text: str):
        super().__init__(input_text)
        self.special_notes = {}
    
    def process(self) -> str:
        """Process Tender Notice document."""
        text = self.raw_text
        
        # Format main header and dates
        text = self._format_header_section(text)
        
        # Process special notes and markers
        text = self._process_special_notes(text)
        
        # Format main sections
        text = self._format_main_sections(text)
        
        # Format rental sections
        text = self._format_rental_sections(text)
        
        # Process location-specific rules
        text = self._format_location_rules(text)
        
        # Format tender details
        text = self._format_tender_details(text)
        
        self.processed_text = text
        return text
    
    def _format_header_section(self, text: str) -> str:
        """Format the main header and tender dates section."""
        dates_pattern = r"\[Opening on .*?\] *\n *\[Closing on .*?\]"
        
        def format_dates(match):
            dates_text = match.group(0)
            return """
## Tender Dates
### Opening
{opening_date}
### Closing
{closing_date}
""".format(
                opening_date=re.search(r"\[Opening on (.*?)\]", dates_text).group(1),
                closing_date=re.search(r"\[Closing on (.*?)\]", dates_text).group(1)
            )
        
        text = re.sub(dates_pattern, format_dates, text)
        text = re.sub(r"TENDER NOTICE", "# TENDER NOTICE\n## August 2024", text)
        return text
    
    def _process_special_notes(self, text: str) -> str:
        """Process and format special notes and markers."""
        text = re.sub(r"\[Important Notes\]:\s*", "\n## Important Notes\n", text)
        
        markers = {
            r"\*\*\*\*": "SPECIAL_NOTE_4",
            r"\*\*\*": "SPECIAL_NOTE_3",
            r"\*\*": "SPECIAL_NOTE_2",
            r"\*": "SPECIAL_NOTE_1",
            r"\+": "HALAL_NOTE",
            r"\^": "INDIAN_CUISINE_NOTE"
        }
        
        for marker, replacement in markers.items():
            notes = re.findall(f"{marker}\\s*(.*?)\\s*(?={marker}|$)", text, re.DOTALL)
            if notes:
                self.special_notes[replacement] = notes
            
            text = re.sub(
                f"{marker}\\s*(.*?)\\s*(?={marker}|$)",
                f"\n> Note ({replacement}): \\1\n",
                text
            )
        
        return text
    
    def _format_main_sections(self, text: str) -> str:
        """Format main document sections."""
        section_patterns = {
            r"Tenders for Rental of \[Cooked Food\] Stalls":
                "\n## Cooked Food Stall Rentals\n",
            r"Tenders for Rental of \[Market\] Stalls":
                "\n## Market Stall Rentals\n",
            r"\[Details of Tender\]":
                "\n## Tender Details\n",
            r"\[Important Notes for All Tenderers\]":
                "\n## General Important Notes\n"
        }
        
        for pattern, replacement in section_patterns.items():
            text = re.sub(pattern, replacement, text)
        
        return text
    
    def _format_rental_sections(self, text: str) -> str:
        """Format rental-specific sections."""
        text = re.sub(
            r"(\n\* Please choose only \[ONE\] type of trade of sale.*?)\n",
            "\n### Trade Type Selection Requirements\\1\n",
            text
        )
        
        text = re.sub(
            r"(\n\* Not for sale of.*?at \[.*?\].*?)\n",
            "\n### Location-Specific Restrictions\\1\n",
            text
        )
        
        return text
    
    def _format_location_rules(self, text: str) -> str:
        """Format location-specific rules."""
        location_pattern = r"\[([^\]]+)\](?=.*?(?:not allowed|not permitted|restricted))"
        
        def format_location(match):
            location = match.group(1)
            return f"\n#### Restrictions for {location}\n"
        
        text = re.sub(location_pattern, format_location, text)
        return text
    
    def _format_tender_details(self, text: str) -> str:
        """Format tender details section."""
        text = re.sub(
            r"Eligibility Criteria:(.*?)(?=-|$)",
            "### Eligibility Requirements\\1",
            text,
            flags=re.DOTALL
        )
        
        text = re.sub(
            r"Tender bids shall be submitted.*?(?=-|$)",
            "### Submission Requirements\\0",
            text,
            flags=re.DOTALL
        )
        
        return text
    
    def get_special_notes(self) -> Dict[str, List[str]]:
        """Return processed special notes."""
        return self.special_notes

class DocumentProcessor:
    """Main document processing coordinator."""
    
    @staticmethod
    def process_docx(file_path: Path) -> str:
        """Convert DOCX to text."""
        doc = docx.Document(file_path)
        return "\n".join([paragraph.text for paragraph in doc.paragraphs])
    
    @staticmethod
    def process_document(file_path: Path) -> ProcessedDocument:
        """Process a document based on its type."""
        file_name = file_path.name.lower()
        
        # Read the file content
        if file_path.suffix == '.docx':
            raw_text = DocumentProcessor.process_docx(file_path)
        else:
            with open(file_path, 'r', encoding='utf-8') as f:
                raw_text = f.read()
        
        # Determine document type and process accordingly
        if "faq" in file_name:
            preprocessor = FAQPreprocessor(raw_text)
            processed_text = preprocessor.process()
            return ProcessedDocument(
                content=processed_text,
                metadata={
                    "source": file_path.name,
                    "type": "faq",
                    "date": "Mar 2024"
                }
            )
        if "terms and conditions" in file_name:
            preprocessor = TenderTermsPreprocessor(raw_text)
            processed_text = preprocessor.process()
            return ProcessedDocument(
                content=processed_text,
                metadata={
                    "source": file_path.name,
                    "type": "terms_and_conditions",
                    "date": "Aug 2024"
                }
            )
        
        elif "tender notice" in file_name:
            preprocessor = TenderNoticePreprocessor(raw_text)
            processed_text = preprocessor.process()
            return ProcessedDocument(
                content=processed_text,
                metadata={
                    "source": file_path.name,
                    "type": "tender_notice",
                    "date": "Aug 2024"
                },
                special_notes=preprocessor.get_special_notes()
            )
        
        else:
            # For other documents, return as-is with basic metadata
            return ProcessedDocument(
                content=raw_text,
                metadata={
                    "source": file_path.name,
                    "type": "general"
                }
            )

@st.cache_data
def load_documents(data_dir: str = DATA_DIR) -> List[Document]:
    """Load and process all documents."""
    documents = []
    print("\nStarting document loading process...")
    
    # Process terms and conditions
    terms_path = Path(data_dir) / "Terms and Conditions of eTender (Aug 2024)_Text file.txt"
    if terms_path.exists():
        print(f"Processing Terms and Conditions from: {terms_path}")
        processed_doc = DocumentProcessor.process_document(terms_path)
        documents.append(Document(
            page_content=processed_doc.content,
            metadata={
                "source": "terms_and_conditions",
                "type": "policy_document",
                "date": "Aug 2024"
            }
        ))
        print("✓ Loaded Terms and Conditions")
    
    # Process tender notice
    notice_path = Path(data_dir) / "Aug 2024 Tender Notice_Text Only.docx"
    if notice_path.exists():
        print(f"Processing Tender Notice from: {notice_path}")
        processed_doc = DocumentProcessor.process_document(notice_path)
        
        # Combine special notes with main content
        if processed_doc.special_notes:
            notes_section = "\n\n## Special Notes and Requirements\n"
            for note_type, notes in processed_doc.special_notes.items():
                notes_section += f"\n### {note_type}\n" + "\n".join(notes)
            
            # Append notes to main content
            processed_doc.content += notes_section
        
        # Add as single document
        documents.append(Document(
            page_content=processed_doc.content,
            metadata={
                "source": "tender_notice",
                "type": "policy_document",
                "date": "Aug 2024"
            }
        ))
        print("✓ Loaded Tender Notice")
    
    # Load AOS data
    json_path = Path(data_dir) / "article_of_sale.json"
    if json_path.exists():
        print(f"Loading AOS data from: {json_path}")
        with open(json_path, 'r', encoding='utf-8') as f:
            articles = json.load(f)
            print(f"Found {len(articles)} articles of sale")
            
            aos_content = "# Guide to Articles of Sale\n\n"
            categories = {}
            for article in articles:
                category = article['Trade Type Category']
                if category not in categories:
                    categories[category] = []
                categories[category].append(article)
            
            for category, items in categories.items():
                aos_content += f"\n## Trade Type: {category}\n"
                for item in items:
                    aos_content += f"\n### {item['Stall Type']}\n"
                    aos_content += f"Article of Sale: {item['Article of Sale']}\n"
                    if item['Remarks']:
                        aos_content += f"Special Notes: {item['Remarks']}\n"
            
            documents.append(Document(
                page_content=aos_content,
                metadata={
                    "source": "articles_of_sale",
                    "type": "reference_guide"
                }
            ))
            print("✓ Loaded Articles of Sale guide")
    
    # Load other text files
    text_files = ["ThingsToNoteWhenTendering.txt"]
    for file in text_files:
        file_path = Path(data_dir) / file
        if file_path.exists():
            print(f"Processing: {file}")
            loader = TextLoader(str(file_path))
            loaded_docs = loader.load()
            for doc in loaded_docs:
                doc.metadata.update({
                    "source": file,
                    "type": "general_info"
                })
                documents.append(doc)
            print(f"✓ Loaded {file}")

    print(f"\nDocument Loading Summary:")
    print(f"- Total documents: {len(documents)}")
    for doc in documents:
        print(f"- {doc.metadata['source']}: {doc.metadata['type']}")
    
    return documents

@st.cache_resource
def create_vector_store(_documents: List[Document]) -> FAISS:
    """Create and cache the vector store."""
    print("\nCreating vector store...")
    
    # Configure text splitter
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1500,
        chunk_overlap=200,
        length_function=len,
        separators=["\n## ", "\n### ", "\n#### ", "\n##### ", "\n\n", "\n", " ", ""]
    )
    
    # Split documents
    texts = text_splitter.split_documents(_documents)
    print(f"Split into {len(texts)} text chunks")
    
    # Create embeddings
    embeddings = OpenAIEmbeddings()
    
    # Create and return vector store
    return FAISS.from_documents(texts, embeddings)

@st.cache_resource
def setup_hawker_guru() -> ConversationalRetrievalChain:
    """Setup the Hawker Guru chatbot."""
    print("\nSetting up Hawker Guru...")
    
    # Load and process documents
    documents = load_documents()
    
    # Create vector store
    vectorstore = create_vector_store(documents)
    
    # Create chat model
    llm = ChatOpenAI(
        temperature=0,
        model_name="gpt-4o-mini"
    )
    
    # Configure custom prompt template properly
    CUSTOM_PROMPT = PromptTemplate(
        input_variables=["context", "chat_history", "question"],
        template="""You are HawkerGuru, an expert assistant for Singapore hawker stall bidding. Use the following context to answer the question at the end. 
        If you don't know the answer, just say that you don't know, don't try to make up an answer. 
        Always try to cite specific sections or documents when providing information.

        Key guidelines:
        1. For general questions about tendering process:
           - First cite relevant sections from Terms and Conditions
           - Then reference any applicable parts of the Tender Notice
           - Explain in clear, simple terms
        
        2. For location-specific questions:
           - First state the general rule
           - Then highlight any special requirements for that location
           - Note any restrictions or special conditions
        
        3. For food type questions:
           - Reference the Articles of Sale guide
           - Note any location-specific restrictions
           - Explain any special requirements (e.g., Halal certification)

        Context: {context}
        
        Chat History: {chat_history}
        
        Question: {question}
        
        Helpful Answer:"""
    )
    
    # Create conversational chain
    qa_chain = ConversationalRetrievalChain.from_llm(
        llm=llm,
        retriever=vectorstore.as_retriever(
            search_kwargs={
                "k": 6,  # Retrieve more context
                "score_threshold": 0.7  # Only include relevant matches
            }
        ),
        return_source_documents=True,
        combine_docs_chain_kwargs={"prompt": CUSTOM_PROMPT},
        verbose=True
    )
    
    print("Hawker Guru setup complete!")
    return qa_chain

class DocumentManager:
    """Manages document operations and updates."""
    
    def __init__(self, data_dir: str = DATA_DIR):
        self.data_dir = Path(data_dir)
        self.processor = DocumentProcessor()
    
    def refresh_documents(self) -> None:
        """Force refresh of document cache."""
        # Clear Streamlit caches
        load_documents.clear()
        create_vector_store.clear()
        setup_hawker_guru.clear()
        
        # Reload documents
        setup_hawker_guru()
    
    def get_document_status(self) -> Dict[str, Dict]:
        """Get status of all documents."""
        status = {}
        for file_path in self.data_dir.glob("*"):
            if file_path.is_file():
                status[file_path.name] = {
                    "last_modified": datetime.fromtimestamp(file_path.stat().st_mtime),
                    "size": file_path.stat().st_size,
                    "type": self._get_document_type(file_path)
                }
        return status
    
    def _get_document_type(self, file_path: Path) -> str:
        """Determine document type based on filename and content."""
        file_name = file_path.name.lower()
        if "terms and conditions" in file_name:
            return "terms_and_conditions"
        elif "tender notice" in file_name:
            return "tender_notice"
        elif file_path.suffix == ".json" and "article" in file_name:
            return "article_of_sale"
        else:
            return "general"

# Initialize document manager if running as main module
if __name__ == "__main__":
    doc_manager = DocumentManager()
    status = doc_manager.get_document_status()
    print("\nDocument Status:")
    for doc, info in status.items():
        print(f"\n{doc}:")
        print(f"  Type: {info['type']}")
        print(f"  Last Modified: {info['last_modified']}")
        print(f"  Size: {info['size']} bytes")