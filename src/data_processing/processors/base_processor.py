import os
from pathlib import Path
from typing import List, Dict, Tuple, Optional, Union
from dataclasses import dataclass
from datetime import datetime
import logging
import re
import json
import docx
import streamlit as st
from dotenv import load_dotenv
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_openai import OpenAIEmbeddings, ChatOpenAI
from langchain_community.vectorstores import FAISS
from langchain.chains import ConversationalRetrievalChain
from langchain.docstore.document import Document
from langchain_community.document_loaders import TextLoader, UnstructuredWordDocumentLoader
from langchain.prompts import PromptTemplate

# Environment setup
if load_dotenv('.env'):
    OPENAI_KEY = os.getenv('OPENAI_API_KEY')
else:
    OPENAI_KEY = st.secrets['OPENAI_API_KEY']

PROJECT_ROOT = Path(__file__).parent.parent.parent
DATA_DIR = os.path.join(PROJECT_ROOT, 'data')

logger = logging.getLogger(__name__)

@dataclass
class ProcessedDocument:
    """Container for processed document content and metadata."""
    content: str
    metadata: Dict
    special_notes: Optional[Dict] = None

class BasePreprocessor:
    """Base class for document preprocessing."""
    
    def __init__(self, input_text: str):
        self.raw_text = input_text
        self.processed_text = ""
        
    def process(self) -> str:
        """Main processing pipeline to be implemented by subclasses."""
        raise NotImplementedError
    
    def get_section_content(self, section_name: str) -> str:
        """Extract content of a specific section."""
        if not self.processed_text:
            self.process()
            
        pattern = fr"## {section_name}.*?(?=## |\Z)"  # Note the 'r' prefix
        match = re.search(pattern, self.processed_text, re.DOTALL)
        return match.group(0) if match else ""
    
    def _basic_cleanup(self, text: str) -> str:
        """Initial cleanup of text."""
        # Remove any null bytes and normalize whitespace
        text = text.replace('\x00', '')
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        
        # Remove image references and other artifacts
        text = re.sub(r'!\[.*?\]\(.*?\)', '', text)
        text = re.sub(r'\{.*?\}', '', text)
        text = re.sub(r'\[.*?\]', '', text)
        
        # Remove empty lines and normalize spacing
        lines = [line.strip() for line in text.split('\n')]
        text = '\n'.join(line for line in lines if line)
        
        return text

class DocumentProcessor:
    """Main document processing coordinator."""
    
    @staticmethod
    def process_docx(file_path: Path) -> str:
        """Convert DOCX to text."""
        doc = docx.Document(file_path)
        return "\n".join([paragraph.text for paragraph in doc.paragraphs])
    
    @staticmethod
    def process_document(file_path: Path) -> ProcessedDocument:
        """Process a document based on its type."""
        # Import processors here to avoid circular imports
        from .faq_processor import FAQPreprocessor
        from .tender_terms_processor import TenderTermsPreprocessor
        from .tender_notice_processor import TenderNoticePreprocessor
        
        file_name = file_path.name.lower()
        
        # Rest of the method stays the same...
        if file_path.suffix == '.docx':
            raw_text = DocumentProcessor.process_docx(file_path)
        else:
            with open(file_path, 'r', encoding='utf-8') as f:
                raw_text = f.read()
                
        # Determine document type and process accordingly
        if "faq" in file_name:
            preprocessor = FAQPreprocessor(raw_text)
            processed_text = preprocessor.process()
            return ProcessedDocument(
                content=processed_text,
                metadata={
                    "source": file_path.name,
                    "type": "faq",
                    "date": "Mar 2024"
                }
            )
        elif "terms and conditions" in file_name:
            preprocessor = TenderTermsPreprocessor(raw_text)
            processed_text = preprocessor.process()
            return ProcessedDocument(
                content=processed_text,
                metadata={
                    "source": file_path.name,
                    "type": "terms_and_conditions",
                    "date": "Aug 2024"
                }
            )
        elif "tender notice" in file_name:
            preprocessor = TenderNoticePreprocessor(raw_text)
            processed_text = preprocessor.process()
            return ProcessedDocument(
                content=processed_text,
                metadata={
                    "source": file_path.name,
                    "type": "tender_notice",
                    "date": "Aug 2024"
                },
                special_notes=preprocessor.get_special_notes()
            )
        else:
            # For other documents, return as-is with basic metadata
            return ProcessedDocument(
                content=raw_text,
                metadata={
                    "source": file_path.name,
                    "type": "general"
                }
            )

@st.cache_data
def load_documents(data_dir: str = DATA_DIR) -> List[Document]:
    """Load and process all documents."""
    documents = []
    print("\nStarting document loading process...")
    
    # Process terms and conditions
    terms_path = Path(data_dir) / "Terms and Conditions of eTender (Aug 2024)_Text file.txt"
    if terms_path.exists():
        print(f"Processing Terms and Conditions from: {terms_path}")
        processed_doc = DocumentProcessor.process_document(terms_path)
        documents.append(Document(
            page_content=processed_doc.content,
            metadata={
                "source": "terms_and_conditions",
                "type": "policy_document",
                "date": "Aug 2024"
            }
        ))
        print("✓ Loaded Terms and Conditions")
    
    # Process tender notice
    notice_path = Path(data_dir) / "Aug 2024 Tender Notice_Text Only.docx"
    if notice_path.exists():
        print(f"Processing Tender Notice from: {notice_path}")
        processed_doc = DocumentProcessor.process_document(notice_path)
        
        # Combine special notes with main content
        if processed_doc.special_notes:
            notes_section = "\n\n## Special Notes and Requirements\n"
            for note_type, notes in processed_doc.special_notes.items():
                notes_section += f"\n### {note_type}\n" + "\n".join(notes)
            
            # Append notes to main content
            processed_doc.content += notes_section
        
        # Add as single document
        documents.append(Document(
            page_content=processed_doc.content,
            metadata={
                "source": "tender_notice",
                "type": "policy_document",
                "date": "Aug 2024"
            }
        ))
        print("✓ Loaded Tender Notice")
    
    # Load AOS data
    json_path = Path(data_dir) / "article_of_sale.json"
    if json_path.exists():
        print(f"Loading AOS data from: {json_path}")
        with open(json_path, 'r', encoding='utf-8') as f:
            articles = json.load(f)
            print(f"Found {len(articles)} articles of sale")
            
            aos_content = "# Guide to Articles of Sale\n\n"
            categories = {}
            for article in articles:
                category = article['Trade Type Category']
                if category not in categories:
                    categories[category] = []
                categories[category].append(article)
            
            for category, items in categories.items():
                aos_content += f"\n## Trade Type: {category}\n"
                for item in items:
                    aos_content += f"\n### {item['Stall Type']}\n"
                    aos_content += f"Article of Sale: {item['Article of Sale']}\n"
                    if item['Remarks']:
                        aos_content += f"Special Notes: {item['Remarks']}\n"
            
            documents.append(Document(
                page_content=aos_content,
                metadata={
                    "source": "articles_of_sale",
                    "type": "reference_guide"
                }
            ))
            print("✓ Loaded Articles of Sale guide")
    
    # Load other text files
    text_files = ["ThingsToNoteWhenTendering.txt"]
    for file in text_files:
        file_path = Path(data_dir) / file
        if file_path.exists():
            print(f"Processing: {file}")
            loader = TextLoader(str(file_path))
            loaded_docs = loader.load()
            for doc in loaded_docs:
                doc.metadata.update({
                    "source": file,
                    "type": "general_info"
                })
                documents.append(doc)
            print(f"✓ Loaded {file}")

    print(f"\nDocument Loading Summary:")
    print(f"- Total documents: {len(documents)}")
    for doc in documents:
        print(f"- {doc.metadata['source']}: {doc.metadata['type']}")
    
    return documents

@st.cache_resource
def create_vector_store(_documents: List[Document]) -> FAISS:
    """Create and cache the vector store."""
    print("\nCreating vector store...")
    
    # Configure text splitter
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1500,
        chunk_overlap=200,
        length_function=len,
        separators=["\n## ", "\n### ", "\n#### ", "\n##### ", "\n\n", "\n", " ", ""]
    )
    
    # Split documents
    texts = text_splitter.split_documents(_documents)
    print(f"Split into {len(texts)} text chunks")
    
    # Create embeddings
    embeddings = OpenAIEmbeddings()
    
    # Create and return vector store
    return FAISS.from_documents(texts, embeddings)

@st.cache_resource
def setup_hawker_guru() -> ConversationalRetrievalChain:
    """Setup the Hawker Guru chatbot."""
    print("\nSetting up Hawker Guru...")
    
    # Load and process documents
    documents = load_documents()
    
    # Create vector store
    vectorstore = create_vector_store(documents)
    
    # Create chat model
    llm = ChatOpenAI(
        temperature=0,
        model_name="gpt-4o-mini"
    )
    
    # Configure custom prompt template properly
    CUSTOM_PROMPT = PromptTemplate(
        input_variables=["context", "chat_history", "question"],
        template="""You are HawkerGuru, an expert assistant for Singapore hawker stall bidding. Use the following context to answer the question at the end. 
        If you don't know the answer, just say that you don't know, don't try to make up an answer. 
        Always try to cite specific sections or documents when providing information.

        Key guidelines:
        1. For general questions about tendering process:
           - First cite relevant sections from Terms and Conditions
           - Then reference any applicable parts of the Tender Notice
           - Explain in clear, simple terms
        
        2. For location-specific questions:
           - First state the general rule
           - Then highlight any special requirements for that location
           - Note any restrictions or special conditions
        
        3. For food type questions:
           - Reference the Articles of Sale guide
           - Note any location-specific restrictions
           - Explain any special requirements (e.g., Halal certification)

        Context: {context}
        
        Chat History: {chat_history}
        
        Question: {question}
        
        Helpful Answer:"""
    )
    
    # Create conversational chain
    qa_chain = ConversationalRetrievalChain.from_llm(
        llm=llm,
        retriever=vectorstore.as_retriever(
            search_kwargs={
                "k": 6,  # Retrieve more context
                "score_threshold": 0.7  # Only include relevant matches
            }
        ),
        return_source_documents=True,
        combine_docs_chain_kwargs={"prompt": CUSTOM_PROMPT},
        verbose=True
    )
    
    print("Hawker Guru setup complete!")
    return qa_chain

class DocumentManager:
    """Manages document operations and updates."""
    
    def __init__(self, data_dir: str = DATA_DIR):
        self.data_dir = Path(data_dir)
        self.processor = DocumentProcessor()
    
    def refresh_documents(self) -> None:
        """Force refresh of document cache."""
        # Clear Streamlit caches
        load_documents.clear()
        create_vector_store.clear()
        setup_hawker_guru.clear()
        
        # Reload documents
        setup_hawker_guru()
    
    def get_document_status(self) -> Dict[str, Dict]:
        """Get status of all documents."""
        status = {}
        for file_path in self.data_dir.glob("*"):
            if file_path.is_file():
                status[file_path.name] = {
                    "last_modified": datetime.fromtimestamp(file_path.stat().st_mtime),
                    "size": file_path.stat().st_size,
                    "type": self._get_document_type(file_path)
                }
        return status
    
    def _get_document_type(self, file_path: Path) -> str:
        """Determine document type based on filename and content."""
        file_name = file_path.name.lower()
        if "terms and conditions" in file_name:
            return "terms_and_conditions"
        elif "tender notice" in file_name:
            return "tender_notice"
        elif file_path.suffix == ".json" and "article" in file_name:
            return "article_of_sale"
        else:
            return "general"

# Initialize document manager if running as main module
if __name__ == "__main__":
    doc_manager = DocumentManager()
    status = doc_manager.get_document_status()
    print("\nDocument Status:")
    for doc, info in status.items():
        print(f"\n{doc}:")
        print(f"  Type: {info['type']}")
        print(f"  Last Modified: {info['last_modified']}")
        print(f"  Size: {info['size']} bytes")